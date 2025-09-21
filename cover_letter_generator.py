import os
import logging
from dotenv import load_dotenv
from langchain_google_genai import ChatGoogleGenerativeAI
from langchain.prompts import PromptTemplate
from docx import Document
from docx.shared import Pt

load_dotenv()
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

COVER_LETTER_FILE = "cover_letter.txt"
COVER_LETTER_DOCX = "cover_letter.docx"

def save_text_to_docx(text, docx_file, title="Cover Letter"):
    doc = Document()
    doc.add_heading(title, 0)

    for line in text.split("\n"):
        if line.strip() == "":
            doc.add_paragraph()
        else:
            p = doc.add_paragraph(line)
            p.style.font.size = Pt(12)

    doc.save(docx_file)
    logger.info(f"Word document saved to {docx_file}")

def generate_cover_letter(resume_text, job_description):
    logger.info("Generating cover letter...")
    
    # Check if API key is available
    api_key = os.getenv("GOOGLE_API")
    if not api_key:
        raise ValueError("Google API key not found. Please set the GOOGLE_API environment variable.")
    
    # Validate inputs
    if not resume_text or not resume_text.strip():
        raise ValueError("Resume text is empty or invalid.")
    
    if not job_description or not job_description.strip():
        raise ValueError("Job description is empty or invalid.")

    llm = ChatGoogleGenerativeAI(
        model="gemini-1.5-flash",
        google_api_key=api_key,
        temperature=0.7
    )

    prompt = PromptTemplate(
        input_variables=["resume", "job_desc"],
        template="""
Using this refined resume:
{resume}

And this job description:
{job_desc}

Generate a professional cover letter tailored for this job.
"""
    )

    try:
        final_prompt = prompt.format(resume=resume_text, job_desc=job_description)
        response = llm.invoke(final_prompt)

        cover_letter = response.content
        
        # Save to files only if not running in Streamlit
        if not os.getenv("STREAMLIT_RUNNING"):
            with open(COVER_LETTER_FILE, "w", encoding="utf-8") as f:
                f.write(cover_letter)
                logger.info(f"Cover letter saved to {COVER_LETTER_FILE}")

            save_text_to_docx(cover_letter, COVER_LETTER_DOCX)
        
        logger.info("Cover letter generation completed successfully.")
        return cover_letter
        
    except Exception as e:
        logger.error(f"Error during cover letter generation: {e}")
        raise
