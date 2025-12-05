import os
import logging
from dotenv import load_dotenv
from langchain_google_genai import ChatGoogleGenerativeAI
from langchain.prompts import PromptTemplate
from docx import Document
from docx.shared import Pt
from app.db_setup.chromadb import get_similar_job_template, store_job_application

load_dotenv()
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

COVER_LETTER_FILE = "cover_letter.txt"
COVER_LETTER_DOCX = "cover_letter.docx"

def save_text_to_docx(text, docx_file, title="Cover Letter"):
    doc = Document()
    doc.add_heading(title, 0)

    for line in text.split("\n"):
        if line.strip() == "":
            doc.add_paragraph()
        else:
            p = doc.add_paragraph(line)
            p.style.font.size = Pt(12)

    doc.save(docx_file)
    logger.info(f"Word document saved to {docx_file}")

def generate_cover_letter(resume_text, job_description):
    logger.info("Generating cover letter...")
    
    api_key = os.getenv("GOOGLE_API")
    if not api_key:
        raise ValueError("Google API key not found.")
    
    if not resume_text or not resume_text.strip():
        raise ValueError("Resume text is empty.")
    
    if not job_description or not job_description.strip():
        raise ValueError("Job description is empty.")

    similar_template = get_similar_job_template(job_description)
    
    if similar_template:
        logger.info(f"Found similar job template! Similarity: {similar_template['similarity']:.2f}")
        
        llm = ChatGoogleGenerativeAI(
            model="gemini-2.5-flash",
            google_api_key=api_key,
            temperature=0.7
        )

        customization_prompt = PromptTemplate(
            input_variables=["template", "resume", "job_desc"],
            template="""
Customize this cover letter template for the candidate's specific resume:

TEMPLATE:
{template}

CANDIDATE RESUME:
{resume}

JOB DESCRIPTION:
{job_desc}

Create a NEW cover letter that:
1. Uses the structure/style of the template
2. Contains information from THIS candidate's resume
3. Addresses THIS job requirements
4. Is completely personalized for this candidate
"""
        )

        final_prompt = customization_prompt.format(
            template=similar_template['cover_letter_template'],
            resume=resume_text,
            job_desc=job_description
        )
        response = llm.invoke(final_prompt)
        cover_letter = response.content
    else:
        llm = ChatGoogleGenerativeAI(
            model="gemini-2.5-flash",
            google_api_key=api_key,
            temperature=0.7
        )

        prompt = PromptTemplate(
            input_variables=["resume", "job_desc"],
            template="""
Using this refined resume:
{resume}

And this job description:
{job_desc}

Generate a professional cover letter tailored for this job.
"""
        )

        final_prompt = prompt.format(resume=resume_text, job_desc=job_description)
        response = llm.invoke(final_prompt)
        cover_letter = response.content

    store_job_application(resume_text, job_description, cover_letter)
    
    if not os.getenv("STREAMLIT_RUNNING"):
        with open(COVER_LETTER_FILE, "w", encoding="utf-8") as f:
            f.write(cover_letter)
        save_text_to_docx(cover_letter, COVER_LETTER_DOCX)
    
    logger.info("Cover letter generated and stored successfully.")
    return cover_letter
