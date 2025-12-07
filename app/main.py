from fastapi import FastAPI, UploadFile, File,Depends
from fastapi.responses import FileResponse
from fastapi.middleware.cors import CORSMiddleware
from db_setup.auth import create_access_token, current_user
from db_setup.database import SessionLocal, User
from service.resume_refiner import refine_resume
from service.cover_letter_generator import generate_cover_letter
from fastapi import HTTPException
from pydantic import BaseModel
import hashlib
import uuid
import uvicorn
import os

app=FastAPI()
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

def hash_password(password: str) -> str:
    return hashlib.sha256(password.encode()).hexdigest()

def verify_password(password: str, hashed: str) -> bool:
    return hashlib.sha256(password.encode()).hexdigest() == hashed

class SignupRequest(BaseModel):
    name: str
    email: str
    password: str

class LoginRequest(BaseModel):
    email: str
    password: str

class CoverLetterRequest(BaseModel):
    resume_text: str
    job_description: str

@app.post("/signup")
def signup(request: SignupRequest):
    db=SessionLocal()
    try:
        if db.query(User).filter(User.email==request.email).first():
            raise HTTPException(400,"User already exists")
        
        hashed=hash_password(request.password)
        user=User(name=request.name, email=request.email, password=hashed)
        db.add(user)
        db.commit()
        return {"message":"User created successfully"}
    finally:
        db.close()

@app.post("/login")
def login(request: LoginRequest):
    db=SessionLocal()
    try:
        user=db.query(User).filter(User.email==request.email).first()
        if not user:
            raise HTTPException(404, "User not found")
        if not verify_password(request.password,user.password):
            raise HTTPException(401,"Wrong password")
        token=create_access_token({"sub":user.email})
        return {"access_token":token, "token_type":"bearer"}
    finally:
        db.close()

@app.post("/refiner")
async def resume_refiner(
    file: UploadFile = File(...),
    current_user: str = Depends(current_user)
):
    temp_path = f"temp_{uuid.uuid4()}.pdf"

    with open(temp_path, "wb") as f:
        f.write(await file.read())

    refined_output = refine_resume(temp_path)

    os.remove(temp_path)
    return refined_output

@app.post("/cover-letter")
async def cover_letter(
    file: UploadFile = File(...),
    job_description: str = File(...),
    current_user: str = Depends(current_user)
    ):
    try:
        temp_pdf = f"temp_{uuid.uuid4()}.pdf"
        with open(temp_pdf, "wb") as f:
            f.write(await file.read())

        refined = refine_resume(temp_pdf)
        refined_resume = refined["refined_resume"]

        cover_letter_text = generate_cover_letter(refined_resume, job_description)

        docx_file = f"cover_letter_{uuid.uuid4()}.docx"

        from docx import Document
        doc = Document()
        doc.add_heading("Cover Letter", 0)

        for line in cover_letter_text.split("\n"):
            doc.add_paragraph(line)

        doc.save(docx_file)

        os.remove(temp_pdf)

        return FileResponse(
            path=docx_file,
            filename="Cover_Letter.docx",
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

    except Exception as e:
        raise HTTPException(500, str(e))


if __name__=="__main__":
    uvicorn.run(app,host="0.0.0.0",port=8000)